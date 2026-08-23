import json
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from lxml import etree

if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

INPUT = r"C:\Users\Administrator\Downloads\BaoCaoDoAn_PRO2192 (AutoRecovered).docx"
OUTPUT = r"D:\FPT\neww\tailieu\BaoCaoDoAn.docx"
LOGO = r"D:\FPT\neww\tailieu\diagrams\logo_dsavisual.png"
TEST_RESULTS = r"D:\FPT\neww\tailieu\test_results.json"

print(f"[Word Fixer] Đang tải tài liệu: {INPUT}")
doc = Document(INPUT)


def fix_font(run, name="Times New Roman", size=12, bold=None):
    run.font.name = name
    run.font.size = Pt(size)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    for attr in [qn("w:ascii"), qn("w:hAnsi"), qn("w:eastAsia"), qn("w:cs")]:
        rFonts.set(attr, name)
    if bold is not None:
        run.bold = bold


def get_size(style_name):
    s = style_name or ""
    if "Heading 1" in s:
        return (14, True)
    if "Heading 2" in s:
        return (13, True)
    if "Heading 3" in s:
        return (12, True)
    if "Caption" in s:
        return (11, None)
    if "TOC" in s:
        return (12, None)
    return (12, None)


# === TASK A: Font chuẩn hóa ===
print("[A] Chuẩn hóa font Times New Roman...")
for para in doc.paragraphs:
    sz, bold = get_size(para.style.name)
    for run in para.runs:
        fix_font(run, size=sz, bold=bold)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    fix_font(run, size=11)

for section in doc.sections:
    for part in [section.header, section.footer]:
        for para in part.paragraphs:
            for run in para.runs:
                fix_font(run, size=10)

# === TASK B: Header — xóa "TÊN ĐỀ TÀI", thêm logo DSA Visual ===
print("[B] Xử lý Header & Logo...")
for section in doc.sections:
    header = section.header
    for para in header.paragraphs:
        runs_to_del = [r._r for r in para.runs if "TÊN ĐỀ TÀI" in r.text]
        for r in runs_to_del:
            try:
                idx = list(para._p).index(r)
                if idx > 0:
                    prev = list(para._p)[idx - 1]
                    if prev.find(qn("w:tab")) is not None:
                        prev.getparent().remove(prev)
                r.getparent().remove(r)
            except Exception:
                pass

        if Path(LOGO).exists():
            try:
                # Kiểm tra tránh chèn lặp logo nếu header đã có
                has_logo = any("graphic" in etree.tostring(r._r).decode("utf-8", errors="ignore") for r in para.runs)
                if not has_logo:
                    run = para.add_run()
                    run.add_picture(LOGO, height=Cm(0.8))
            except Exception as e:
                print(f"  Logo insert info: {e}")

# === TASK C: Xóa song ngữ tiêu đề ===
print("[C] Làm sạch tiêu đề song ngữ...")
BILINGUAL = {
    "KHẢO SÁT – SURVEY": "KHẢO SÁT",
    "KHẢO SÁT - SURVEY": "KHẢO SÁT",
    "PHÂN TÍCH - ANALYSIS": "PHÂN TÍCH",
    "PHÂN TÍCH – ANALYSIS": "PHÂN TÍCH",
    "THIẾT KẾ - DESIGN": "THIẾT KẾ",
    "THIẾT KẾ – DESIGN": "THIẾT KẾ",
    "THỰC HIỆN – IMPLEMENT": "THỰC HIỆN",
    "THỰC HIỆN - IMPLEMENT": "THỰC HIỆN",
    "KIỂM THỬ - TESTING": "KIỂM THỬ",
    "KIỂM THỬ – TESTING": "KIỂM THỬ",
}
for para in doc.paragraphs:
    for old, new in BILINGUAL.items():
        if old in para.text and para.runs:
            para.runs[0].text = para.text.replace(old, new)
            for r in para.runs[1:]:
                r.text = ""

# === TASK D: Xóa ngoặc trong Heading ===
print("[D] Xóa chú thích ngoặc đơn trong Headings...")
for para in doc.paragraphs:
    if para.style.name.startswith("Heading") and para.runs:
        cleaned = re.sub(r"\s*\([^)]+\)\s*$", "", para.text).strip()
        if cleaned != para.text:
            para.runs[0].text = cleaned
            for r in para.runs[1:]:
                r.text = ""

# === TASK E: Bảng auto-fit + nvarchar optimize ===
print("[E] Tối ưu hóa bảng và kiểu dữ liệu...")
for table in doc.tables:
    table.autofit = True
    try:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    except Exception:
        pass
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    fix_font(run, size=11)
                    if "nvarchar(256)" in run.text:
                        run.text = run.text.replace("nvarchar(256)", "nvarchar(100)")
                    if "nvarchar(512)" in run.text:
                        run.text = run.text.replace("nvarchar(512)", "nvarchar(255)")

# === TASK F: Thêm bảng test case Section 6 ===
print("[F] Chèn bảng 133 Test Cases vào Phần Kiểm Thử...")
if Path(TEST_RESULTS).exists():
    with open(TEST_RESULTS, encoding="utf-8") as f:
        data = json.load(f)

    target_para = None
    for para in doc.paragraphs:
        if "KIỂM THỬ" in para.text.upper() and para.style.name.startswith("Heading"):
            target_para = para

    if target_para is not None:
        headers = ["TC-ID", "Màn hình", "Thao tác", "Đầu vào", "Kết quả mong đợi", "Kết quả thực tế", "Kết quả"]
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        table.autofit = True

        hdr = table.rows[0]
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = h
            if cell.paragraphs[0].runs:
                for run in cell.paragraphs[0].runs:
                    fix_font(run, size=10, bold=True)

        for tc in data["testCases"]:
            row = table.add_row()
            vals = [
                tc.get("id", ""),
                tc.get("screen", ""),
                tc.get("action", ""),
                tc.get("input", ""),
                tc.get("expected", ""),
                tc.get("actual", ""),
                tc.get("status", "PASS"),
            ]
            for i, v in enumerate(vals):
                row.cells[i].text = str(v)[:250]
                if row.cells[i].paragraphs[0].runs:
                    for run in row.cells[i].paragraphs[0].runs:
                        fix_font(run, size=9)

        target_para._p.addnext(table._tbl)
        print(f"  ✓ Đã chèn thành công bảng với {len(data['testCases'])} dòng test case")

# === TASK J: Xóa TOC Error entries ===
print("[J] Loại bỏ các dòng TOC bị lỗi...")
for para in list(doc.paragraphs):
    if "Error! Bookmark not defined" in para.text:
        try:
            p = para._p
            p.getparent().remove(p)
        except Exception:
            pass

# === TASK K: Xóa Page Break cứng ở đầu tài liệu ===
print("[K] Dọn dẹp ngắt trang cứng đầu tài liệu...")
for para in doc.paragraphs[:80]:
    for run in para.runs:
        for br in run._r.findall(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                br.getparent().remove(br)

# === SAVE ===
print(f"[SAVE] Lưu tài liệu chuẩn hóa ra: {OUTPUT}")
doc.save(OUTPUT)
print("Hoàn tất xử lý tài liệu Word thành công! ✓")
