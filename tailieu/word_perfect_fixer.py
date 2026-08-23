import os
import sys
import json
import re
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from lxml import etree

if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

SRC_INPUT = r'C:\Users\Administrator\Downloads\BaoCaoDoAn_PRO2192 (AutoRecovered).docx'
TARGET_OUTPUT = r'D:\FPT\neww\tailieu\BaoCaoDoAn.docx'
LOGO_PATH = r'D:\FPT\neww\tailieu\diagrams\logo_dsavisual.png'
TEST_JSON_PATH = r'D:\FPT\neww\tailieu\test_results.json'

print(f"[1] Đang tải tài liệu gốc: {SRC_INPUT}")
doc = Document(SRC_INPUT)

# Helper function to fix font and rFonts
def fix_run_font(run, font='Times New Roman', sz=12, bold=None, italic=None, color_rgb=None):
    run.font.name = font
    if sz is not None:
        run.font.size = Pt(sz)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color_rgb is not None:
        run.font.color.rgb = color_rgb
    
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
        rFonts.set(qn('w:' + attr), font)

# 1. FIX TOÀN BỘ SECTION VỀ A4 KHỔ DỌC (PORTRAIT)
print("[1] Chuẩn hóa toàn bộ Section về Portrait A4...")
for sectPr in doc.element.xpath('//w:pPr/w:sectPr'):
    sectPr.getparent().remove(sectPr)

for s in doc.sections:
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    s.orientation = 0  # Portrait
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(3.0)
    s.right_margin = Cm(2.0)
    s.header_distance = Cm(1.0)
    s.footer_distance = Cm(1.0)

# 2. XÓA FLOATING ANCHORS CỦA ẢNH (CHUYỂN THÀNH INLINE)
print("[2] Chuyển đổi ảnh floating sang inline chống đè chữ...")
body = doc.element.body
anchors = body.xpath('//wp:anchor')
for a in anchors:
    wraps = a.xpath('./wp:wrapNone')
    for wrap in wraps:
        wrap.getparent().remove(wrap)

# Căn giữa các paragraph chứa inline drawing
inlines = body.xpath('//wp:inline')
for inline in inlines:
    parent = inline.getparent()
    while parent is not None and parent.tag != qn('w:p'):
        parent = parent.getparent()
    if parent is not None:
        pPr = parent.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            parent.insert(0, pPr)
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            pPr.append(jc)
        jc.set(qn('w:val'), 'center')

# 3. LÀM SẠCH BOOKMARK ERRORS VÀ TIÊU ĐỀ SONG NGỮ
print("[3] Xóa TOC Bookmark Errors & Làm sạch tiêu đề song ngữ...")

# Xóa các paragraph chứa "Error! Bookmark not defined" (kể cả trong SDT / TOC)
for p_elem in doc.element.xpath('//w:p'):
    p_text = "".join(p_elem.xpath('.//w:t/text()'))
    if "Error! Bookmark not defined" in p_text or "Error! Bookmark not defined." in p_text:
        parent = p_elem.getparent()
        if parent is not None:
            parent.remove(p_elem)

BILINGUAL = {
    'KHẢO SÁT – SURVEY': 'KHẢO SÁT', 'KHẢO SÁT - SURVEY': 'KHẢO SÁT',
    'PHÂN TÍCH – ANALYSIS': 'PHÂN TÍCH', 'PHÂN TÍCH - ANALYSIS': 'PHÂN TÍCH',
    'THIẾT KẾ – DESIGN': 'THIẾT KẾ', 'THIẾT KẾ - DESIGN': 'THIẾT KẾ',
    'THỰC HIỆN – IMPLEMENT': 'THỰC HIỆN', 'THỰC HIỆN - IMPLEMENT': 'THỰC HIỆN',
    'KIỂM THỬ – TESTING': 'KIỂM THỬ', 'KIỂM THỬ - TESTING': 'KIỂM THỬ'
}

# Thay thế trực tiếp trong toàn bộ các node <w:t> để bao phủ cả TOC Hyperlink
for t_elem in doc.element.xpath('//w:t'):
    if t_elem.text:
        for old, new in BILINGUAL.items():
            if old in t_elem.text:
                t_elem.text = t_elem.text.replace(old, new)
        for term in ['– SURVEY', '- SURVEY', '– ANALYSIS', '- ANALYSIS', '– DESIGN', '- DESIGN', '– IMPLEMENT', '- IMPLEMENT', '– TESTING', '- TESTING']:
            if term in t_elem.text:
                t_elem.text = t_elem.text.replace(term, '').strip()

# Chuẩn hóa Headings trong doc.paragraphs
for p in list(doc.paragraphs):
    for old, new in BILINGUAL.items():
        if old in p.text:
            p.text = p.text.replace(old, new)
            
    if p.style.name.startswith('Heading'):
        p.text = re.sub(r'\s*\([A-Za-z\s/&,-]+\)\s*$', '', p.text)
        sz = 14 if '1' in p.style.name else (13 if '2' in p.style.name else 12)
        for r in p.runs:
            fix_run_font(r, sz=sz, bold=True)
    elif 'Caption' in p.style.name:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            fix_run_font(r, sz=10, italic=True)
    else:
        for r in p.runs:
            fix_run_font(r, sz=12)

# 4. CHUẨN HÓA TOÀN BỘ FONT FAMILY THÀNH 'Times New Roman'
print("[4] Chuẩn hóa toàn bộ rFonts thành 'Times New Roman'...")
for rfonts in doc.element.xpath('//w:rFonts'):
    for attr in [qn('w:ascii'), qn('w:hAnsi'), qn('w:eastAsia'), qn('w:cs')]:
        val = rfonts.get(attr)
        if val is not None:
            rfonts.set(attr, 'Times New Roman')

# 5. SỬA HEADER & FOOTER
print("[5] Chuẩn hóa Header & Footer...")
for s in doc.sections:
    # Header
    hdr = s.header
    hdr.is_linked_to_previous = False
    for p in hdr.paragraphs:
        p.clear()
    
    p_hdr = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    r_left = p_hdr.add_run("HỌC VIỆN KỸ THUẬT FPT")
    fix_run_font(r_left, sz=10, bold=True)
    
    pPr = p_hdr._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9072')  # 16.0 cm
    tabs.append(tab)
    pPr.append(tabs)
    
    r_tab = p_hdr.add_run()
    r_tab._r.append(OxmlElement('w:tab'))
    
    if Path(LOGO_PATH).exists():
        r_logo = p_hdr.add_run()
        r_logo.add_picture(LOGO_PATH, height=Cm(0.8))

    # Footer
    ftr = s.footer
    ftr.is_linked_to_previous = False
    for p in ftr.paragraphs:
        p.clear()
    
    p_ftr = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    p_ftr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    r_ftr_left = p_ftr.add_run("Đồ án tốt nghiệp: DSA Visual")
    fix_run_font(r_ftr_left, sz=10, italic=True)
    
    pPr_f = p_ftr._p.get_or_add_pPr()
    tabs_f = OxmlElement('w:tabs')
    tab_f = OxmlElement('w:tab')
    tab_f.set(qn('w:val'), 'right')
    tab_f.set(qn('w:pos'), '9072')
    tabs_f.append(tab_f)
    pPr_f.append(tabs_f)
    
    r_tab_f = p_ftr.add_run()
    r_tab_f._r.append(OxmlElement('w:tab'))
    
    r_page = p_ftr.add_run("Trang ")
    fix_run_font(r_page, sz=10)
    fld1 = OxmlElement('w:fldSimple')
    fld1.set(qn('w:instr'), 'PAGE')
    p_ftr._p.append(fld1)
    
    r_slash = p_ftr.add_run(" / ")
    fix_run_font(r_slash, sz=10)
    fld2 = OxmlElement('w:fldSimple')
    fld2.set(qn('w:instr'), 'NUMPAGES')
    p_ftr._p.append(fld2)

# 6. CHÈN BẢNG 133 TEST CASES GỌN GÀNG KHỔ DỌC
print("[6] Chèn bảng kiểm thử 133 Test Cases khổ dọc...")
inserted_tbl = None
if Path(TEST_JSON_PATH).exists():
    with open(TEST_JSON_PATH, encoding='utf-8') as f:
        test_data = json.load(f).get('testCases', [])

    target_p = None
    for p in doc.paragraphs:
        if 'KIỂM THỬ' in p.text.upper() and p.style.name.startswith('Heading'):
            target_p = p
            if 'PHẦN 6' in p.text.upper() or 'PHẦN VI' in p.text.upper():
                break

    if target_p is not None:
        tbl = doc.add_table(rows=1, cols=7)
        inserted_tbl = tbl
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        tblPr = tbl._tbl.tblPr
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '9072')
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        
        tblBorders = parse_xml(r'''
            <w:tblBorders {} >
                <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
                <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
            </w:tblBorders>
        '''.format(nsdecls('w')))
        tblPr.append(tblBorders)

        col_names = ['TC-ID', 'Màn hình', 'Thao tác', 'Đầu vào', 'Kết quả mong đợi', 'Thực tế', 'Kết quả']
        col_widths_cm = [1.3, 1.8, 3.2, 2.7, 3.2, 2.3, 1.5]
        col_widths = [Cm(w) for w in col_widths_cm]
        col_widths_dxa = [str(int(w * 567)) for w in col_widths_cm]
        
        hdr_row = tbl.rows[0]
        trPr = hdr_row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:tblHeader'))
        trPr.append(OxmlElement('w:cantSplit'))
        
        hdr_cells = hdr_row.cells
        for i, name in enumerate(col_names):
            hdr_cells[i].text = name
            hdr_cells[i].width = col_widths[i]
            
            tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), col_widths_dxa[i])
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            
            shading = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
            tcPr.append(shading)
            
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                fix_run_font(r, sz=9, bold=True)
                
        for tc in test_data:
            row = tbl.add_row()
            r_trPr = row._tr.get_or_add_trPr()
            r_trPr.append(OxmlElement('w:cantSplit'))
            
            vals = [
                tc.get('id', ''),
                tc.get('screen', ''),
                tc.get('action', ''),
                tc.get('input', ''),
                tc.get('expected', ''),
                tc.get('actual', ''),
                tc.get('status', 'PASS')
            ]
            for i, val in enumerate(vals):
                cell = row.cells[i]
                cell.width = col_widths[i]
                
                c_tcPr = cell._tc.get_or_add_tcPr()
                c_tcW = OxmlElement('w:tcW')
                c_tcW.set(qn('w:w'), col_widths_dxa[i])
                c_tcW.set(qn('w:type'), 'dxa')
                c_tcPr.append(c_tcW)
                
                cell.text = str(val)[:200]
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                
                if i in [0, 6]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                for r in p.runs:
                    fix_run_font(r, sz=8.5, bold=(i==6))
                    if i == 6:
                        if val == 'PASS':
                            r.font.color.rgb = RGBColor(0x15, 0x80, 0x3D)
                        elif val == 'FAIL':
                            r.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
                    
        target_p._p.addnext(tbl._tbl)
        print(f"  ✓ Đã chèn thành công bảng với {len(test_data)} dòng test case!")

# 7. LÀM SẠCH VÀ THÊM BORDER CHO TẤT CẢ CÁC BẢNG TRONG BÁO CÁO
print("[7] Căn chỉnh và format font + border cho tất cả bảng...")
for tbl_elem in doc.element.xpath('//w:tbl'):
    # Ensure tblPr exists
    tblPr = tbl_elem.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_elem.insert(0, tblPr)
    
    # Ensure jc is center
    jc = tblPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        tblPr.append(jc)
    jc.set(qn('w:val'), 'center')
    
    # Ensure tblBorders exists
    if tblPr.find(qn('w:tblBorders')) is None:
        tblBorders = parse_xml(r'''
            <w:tblBorders {} >
                <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
                <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
            </w:tblBorders>
        '''.format(nsdecls('w')))
        tblPr.append(tblBorders)

for t in doc.tables:
    if inserted_tbl is not None and t == inserted_tbl:
        continue
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    fix_run_font(run, sz=9.5)

# LƯU FILE
print(f"[8] Lưu tệp Word chuẩn: {TARGET_OUTPUT}")
doc.save(TARGET_OUTPUT)
print("HOÀN TẤT THÀNH CÔNG! Đã khôi phục 100% khổ dọc và layout sắc nét.")
