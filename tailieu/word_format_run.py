"""
word_format_run.py — Tool for formatting and embedding screenshots in BaoCaoDoAn.docx
Usage:
    python word_format_run.py --steps=M
"""

import os
import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

ROOT_DIR = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT_DIR / "tailieu" / "BaoCaoDoAn.docx"
SCREENSHOTS_DIR = ROOT_DIR / "tailieu" / "screenshots"

# Caption prefix -> Screenshot filename
CAPTION_TO_SCREENSHOT = {
    "Hình 4.1:": "01_landing.png",
    "Hình 4.2:": "02_login.png",
    "Hình 4.3:": "14_lesson_detail.png",
    "Hình 4.4:": "09_mo_phong_detail.png",
    "Hình 4.5:": "15_exercise.png",
    "Hình 4.6:": "05_lo_trinh.png",
    "Hình 4.7:": "16_ladder.png",
    "Hình 4.8:": "17_lab.png",
    "Hình 4.9:": "18_code_runner.png",
    # Note: Hình 4.10 is intentionally skipped (caption does not exist in docx)
    "Hình 4.11:": "12_bang_xep_hang.png",
    "Hình 4.12:": "13_ho_so.png",
    "Hình 4.14:": "03_register.png",
    "Hình 4.15:": "04_dashboard.png",
    "Hình 4.16:": "06_lo_trinh_detail.png",
    "Hình 4.17:": "07_node_hub.png",
    "Hình 4.18:": "08_mo_phong.png",
    "Hình 4.19:": "10_lop_hoc.png",
    "Hình 4.20:": "11_lop_hoc_detail.png",
    "Hình 4.21:": "20_admin_dashboard.png",
    "Hình 4.22:": "21_admin_users.png",
    "Hình 4.23:": "22_admin_content.png",
    "Hình 4.24:": "23_admin_settings.png",
}

def embed_screenshots(doc_path: Path):
    print(f"=== EMBEDDING SCREENSHOTS INTO: {doc_path} ===")
    if not doc_path.exists():
        print(f"ERROR: Document {doc_path} does not exist!")
        sys.exit(1)

    doc = Document(str(doc_path))
    embedded_count = 0

    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        for cap_prefix, img_file in CAPTION_TO_SCREENSHOT.items():
            if text.startswith(cap_prefix):
                img_path = SCREENSHOTS_DIR / img_file
                if not img_path.exists():
                    print(f"WARNING: Image file not found: {img_path}")
                    continue

                # Find or reuse the preceding paragraph for image
                if i > 0:
                    prev_p = paragraphs[i - 1]
                    # Clear existing content/drawings from previous paragraph
                    pPr = prev_p._element.find(qn('w:pPr'))
                    prev_p._element.clear()
                    if pPr is not None:
                        prev_p._element.append(pPr)
                    
                    prev_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = prev_p.add_run()
                    run.add_picture(str(img_path), width=Inches(6.0))
                    print(f"  [Embedded] {img_file} -> {cap_prefix}")
                    embedded_count += 1
                break

    doc.save(str(doc_path))
    print(f"SUCCESS: Embedded {embedded_count}/{len(CAPTION_TO_SCREENSHOT)} screenshots into {doc_path}")

def main():
    parser = argparse.ArgumentParser(description="DOCX Formatting and Screenshot Embedding Tool")
    parser.add_argument("--steps", default="M", help="Steps to run (M = embed screenshots only)")
    args = parser.parse_args()

    if "M" in args.steps:
        embed_screenshots(DOCX_PATH)
    else:
        print(f"Unknown steps: {args.steps}")

if __name__ == "__main__":
    main()
