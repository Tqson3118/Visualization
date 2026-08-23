"""
clean_benchmark_docx.py — Remove all Benchmark-related entries from BaoCaoDoAn.docx
"""

import sys
from pathlib import Path
from docx import Document

if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

ROOT_DIR = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT_DIR / "tailieu" / "BaoCaoDoAn.docx"

def remove_row(table, row_idx):
    tr = table.rows[row_idx]._tr
    tr.getparent().remove(tr)

def clean_docx():
    print(f"=== CLEANING BENCHMARK FROM {DOCX_PATH} ===")
    doc = Document(str(DOCX_PATH))

    # 1. Table 2: So sánh hệ thống (Row 8: 'So sánh đo thật với lý thuyết')
    t2 = doc.tables[2]
    for idx, r in enumerate(list(t2.rows)):
        cells_text = [c.text.strip() for c in r.cells]
        if 'So sánh đo thật với lý thuyết' in cells_text[0]:
            print(f"  [Table 2] Removing row: {cells_text}")
            remove_row(t2, idx)
            break

    # 2. Table 3: Phân loại mô phỏng (Row 3: 'Danh mục mô phỏng, hiển thị 3 vùng đồng bộ, điều khiển, Benchmark Lab')
    t3 = doc.tables[3]
    for r in t3.rows:
        if 'Benchmark Lab' in r.cells[1].text:
            r.cells[1].text = 'Danh mục mô phỏng, hiển thị 3 vùng đồng bộ, điều khiển giải thuật'
            r.cells[2].text = 'FR-3.1 đến FR-3.19'
            print("  [Table 3] Updated row text to remove Benchmark Lab")

    # 3. Table 4: Kế hoạch Sprint (Row 9: S9)
    t4 = doc.tables[4]
    for r in t4.rows:
        if 'S9' in r.cells[0].text:
            r.cells[2].text = 'Premium, Lớp học phần'
            r.cells[3].text = 'Gems Shop, Premium mô phỏng, lớp học phần'
            print("  [Table 4] Updated S9 sprint row to remove Benchmark")

    # 4. Table 7: Bảng Use Case (UC-28)
    t7 = doc.tables[7]
    for idx, r in enumerate(list(t7.rows)):
        if len(r.cells) > 0 and r.cells[0].text.strip() == 'UC-28':
            print(f"  [Table 7] Removing UC-28 row: {[c.text.strip() for c in r.cells]}")
            remove_row(t7, idx)
            break

    # 5. Table 10: Bảng Yêu cầu chức năng (FR-3.20, FR-3.20b)
    t10 = doc.tables[10]
    indices_to_remove = []
    for idx, r in enumerate(t10.rows):
        if len(r.cells) > 0 and r.cells[0].text.strip() in ['FR-3.20', 'FR-3.20b']:
            indices_to_remove.append(idx)
    
    for idx in reversed(indices_to_remove):
        print(f"  [Table 10] Removing FR row at index {idx}: {[c.text.strip() for c in t10.rows[idx].cells]}")
        remove_row(t10, idx)

    # 6. Table 12: Phân loại màn hình
    t12 = doc.tables[12]
    for r in t12.rows:
        if '17 Benchmark Lab' in r.cells[1].text:
            r.cells[1].text = '05 Simulator, 33 Khám phá'
            r.cells[2].text = '2'
            print("  [Table 12] Updated screen catalog row")

    # 7. Table 52: Bảng kết quả kiểm thử Test Cases
    t52 = doc.tables[52]
    tc_benchmark_indices = []
    for idx, r in enumerate(t52.rows):
        if len(r.cells) > 0 and r.cells[0].text.strip() in ['TC-125', 'TC-126', 'TC-127']:
            tc_benchmark_indices.append(idx)

    for idx in reversed(tc_benchmark_indices):
        print(f"  [Table 52] Removing Benchmark TC row at {idx}: {[c.text.strip() for c in t52.rows[idx].cells]}")
        remove_row(t52, idx)

    # Renumber remaining TC rows in Table 52 (from TC-128..TC-133 to TC-125..TC-130)
    current_tc = 1
    for r in t52.rows[1:]: # Skip header
        r.cells[0].text = f"TC-{current_tc:02d}"
        current_tc += 1
    print(f"  [Table 52] Renumbered {len(t52.rows)-1} TC rows to TC-01..TC-{current_tc-1:02d}")

    doc.save(str(DOCX_PATH))
    print(f"SUCCESS: Saved cleaned document to {DOCX_PATH}")

if __name__ == "__main__":
    clean_docx()
