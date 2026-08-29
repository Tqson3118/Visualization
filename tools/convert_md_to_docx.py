import re
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_cell_borders(cell, color="CCCCCC", sz="4", val="single"):
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)

def clean_inline_math(text):
    text = text.replace(r'\rightarrow', '→')
    text = text.replace(r'\leftarrow', '←')
    text = text.replace(r'\ge', '≥')
    text = text.replace(r'\le', '≤')
    text = text.replace(r'\blacktriangleright\mid', '▶|')
    text = text.replace(r'\mid\blacktriangleleft', '|◀')
    text = text.replace(r'\blacktriangleright', '▶')
    text = text.replace(r'\blacktriangleleft', '◀')
    text = re.sub(r'\$(.*?)\$', r'\1', text)
    return text

def add_formatted_text(p, text):
    text = clean_inline_math(text)
    # Split text by bold (**), inline code (`), italic (*)
    # Pattern: (\*\*.*?\*\*|`.*?`|\*.*?\*)
    tokens = re.split(r'(\*\*.*?\*\*|`.*?`|\*.*?\*)', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**') and len(token) >= 4:
            run = p.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('`') and token.endswith('`') and len(token) >= 2:
            run = p.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(180, 40, 40)
        elif token.startswith('*') and token.endswith('*') and len(token) >= 2:
            run = p.add_run(token[1:-1])
            run.italic = True
        else:
            p.add_run(token)

def convert_md_file_to_docx(md_path, docx_path):
    print(f"Converting {md_path} -> {docx_path}...")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = docx.Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Base styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(40, 40, 40)

    in_code_block = False
    code_block_lines = []
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        
        # Filter out separator rows (|---|---|)
        valid_rows = []
        for r in table_rows:
            if re.match(r'^\s*\|?\s*:?-+:?\s*(\|?\s*:?-+:?\s*)+\|?\s*$', r):
                continue
            cells = [c.strip() for c in r.strip().strip('|').split('|')]
            if cells and any(cells):
                valid_rows.append(cells)
        
        if not valid_rows:
            in_table = False
            table_rows = []
            return

        cols_count = max(len(r) for r in valid_rows)
        table = doc.add_table(rows=len(valid_rows), cols=cols_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row_idx, row_data in enumerate(valid_rows):
            is_header = (row_idx == 0)
            row = table.rows[row_idx]
            for col_idx in range(cols_count):
                cell = row.cells[col_idx]
                cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
                
                set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                set_cell_borders(cell, color="CBD5E0", sz="4")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.1

                if is_header:
                    set_cell_background(cell, "2B6CB0")
                    run_tokens = re.split(r'(\*\*.*?\*\*|`.*?`)', clean_inline_math(cell_text))
                    for t in run_tokens:
                        if not t: continue
                        if t.startswith('**') and t.endswith('**'):
                            r = p.add_run(t[2:-2])
                        elif t.startswith('`') and t.endswith('`'):
                            r = p.add_run(t[1:-1])
                        else:
                            r = p.add_run(t)
                        r.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.font.size = Pt(10)
                else:
                    if row_idx % 2 == 1:
                        set_cell_background(cell, "F7FAFC")
                    else:
                        set_cell_background(cell, "FFFFFF")
                    add_formatted_text(p, cell_text)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        in_table = False
        table_rows = []

    def flush_code_block():
        nonlocal in_code_block, code_block_lines
        if not code_block_lines:
            in_code_block = False
            return
        
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "EDF2F7")
        set_cell_borders(cell, color="CBD5E0", sz="6")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        
        code_text = "".join(code_block_lines)
        run = p.add_run(code_text.rstrip('\n'))
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(45, 55, 72)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        in_code_block = False
        code_block_lines = []

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle code fence ```
        if stripped.startswith('```'):
            if in_table:
                flush_table()
            if in_code_block:
                flush_code_block()
            else:
                in_code_block = True
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Handle Table
        if '|' in stripped and (stripped.startswith('|') or stripped.endswith('|')):
            in_table = True
            table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            flush_table()

        # Empty line
        if not stripped:
            i += 1
            continue

        # Horizontal Rule
        if stripped in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CBD5E0"/></w:pBdr>')
            p._element.get_or_add_pPr().append(p_border)
            i += 1
            continue

        # Headings
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(clean_inline_math(stripped[2:]))
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(26, 54, 93) # Navy
            i += 1
            continue

        if stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(clean_inline_math(stripped[3:]))
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(43, 108, 176) # Slate Blue
            i += 1
            continue

        if stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(clean_inline_math(stripped[4:]))
            run.bold = True
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(45, 55, 72)
            i += 1
            continue

        if stripped.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(clean_inline_math(stripped[5:]))
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(74, 85, 104)
            i += 1
            continue

        # Blockquote >
        if stripped.startswith('> '):
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = tbl.cell(0, 0)
            set_cell_background(cell, "F7FAFC")
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            
            # left border only
            tcPr = cell._element.get_or_add_tcPr()
            borders = parse_xml(f'''
                <w:tcBorders {nsdecls("w")}>
                    <w:left w:val="single" w:sz="18" w:space="0" w:color="3182CE"/>
                    <w:top w:val="none"/>
                    <w:bottom w:val="none"/>
                    <w:right w:val="none"/>
                </w:tcBorders>
            ''')
            tcPr.append(borders)

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, stripped[2:])
            i += 1
            continue

        # Bullet List (- or *)
        if re.match(r'^\s*[-*]\s+', line):
            indent_level = len(re.match(r'^\s*', line).group(0)) // 2
            item_text = re.sub(r'^\s*[-*]\s+', '', line).strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            p.paragraph_format.space_before = Pt(1.5)
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, item_text)
            i += 1
            continue

        # Numbered List
        if re.match(r'^\s*\d+\.\s+', line):
            indent_level = len(re.match(r'^\s*', line).group(0)) // 2
            item_text = re.sub(r'^\s*\d+\.\s+', '', line).strip()
            num_match = re.match(r'^\s*(\d+\.)\s+', line).group(1)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            
            run_num = p.add_run(num_match + " ")
            run_num.bold = True
            run_num.font.color.rgb = RGBColor(43, 108, 176)
            add_formatted_text(p, item_text)
            i += 1
            continue

        # Normal Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        add_formatted_text(p, stripped)
        i += 1

    if in_table:
        flush_table()
    if in_code_block:
        flush_code_block()

    doc.save(docx_path)
    print(f"Saved {docx_path} successfully!")

if __name__ == '__main__':
    base_dir = r'D:\FPT\neww\docs\test-bank'
    convert_md_file_to_docx(
        os.path.join(base_dir, 'DE_THI_KIEN_THUC_CODEBASE_DSA.md'),
        os.path.join(base_dir, 'DE_THI_KIEN_THUC_CODEBASE_DSA.docx')
    )
    convert_md_file_to_docx(
        os.path.join(base_dir, 'DAP_AN_VA_GIAI_THICH_CHI_TIET.md'),
        os.path.join(base_dir, 'DAP_AN_VA_GIAI_THICH_CHI_TIET.docx')
    )
